from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "Reports"
FIGURES = REPORTS / "figures"
TABLES = REPORTS / "tables"
SRC = ROOT / "src"

ORIGINAL_DOCX = REPORTS / "Project Report  Optimizing Retail Stock Levels Through AI-Powered Demand Prediction.docx"
OUTPUT_DOCX = REPORTS / "Enhanced_Project_Report.docx"


ACCENT = RGBColor(31, 78, 121)
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F3F6F8"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table, header_fill: str = LIGHT_BLUE) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cell, header_fill if r_idx == 0 else WHITE)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9 if r_idx else 9.5)
                    run.font.bold = r_idx == 0


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    r.font.color.rgb = ACCENT
    p.add_run("\n")
    body_run = p.add_run(body)
    body_run.font.name = "Arial"
    body_run.font.size = Pt(10)
    doc.add_paragraph()


def add_csv_table(
    doc: Document,
    csv_path: Path,
    title: str,
    max_rows: int | None = None,
    round_numeric: bool = False,
) -> None:
    if not csv_path.exists():
        return
    with csv_path.open(newline="", encoding="utf-8") as fh:
        rows = list(csv.reader(fh))
    if not rows:
        return
    if max_rows is not None:
        rows = [rows[0], *rows[1 : max_rows + 1]]

    doc.add_heading(title, level=3)
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            value = value.strip()
            if round_numeric and r_idx > 0:
                try:
                    value = f"{float(value):,.2f}"
                except ValueError:
                    pass
            table.cell(r_idx, c_idx).text = value
    style_table(table)
    add_caption(doc, f"{title} (generated from {csv_path.name})")


def add_picture(doc: Document, image_path: Path, title: str, width: float = 6.2) -> None:
    if not image_path.exists():
        return
    doc.add_heading(title, level=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    add_caption(doc, f"{title} ({image_path.name})")


def read_code_lines(path: Path, start: int, end: int) -> str:
    lines = path.read_text(encoding="utf-8").splitlines()
    selected = lines[start - 1 : end]
    return "\n".join(f"{idx + start:03d}: {line}" for idx, line in enumerate(selected))


def add_code_block(doc: Document, title: str, code: str) -> None:
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    set_cell_margins(cell, top=100, bottom=100, start=120, end=120)
    p = cell.paragraphs[0]
    for line_no, line in enumerate(code.splitlines()):
        if line_no:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8)


def add_static_contents(doc: Document) -> None:
    doc.add_heading("Contents", level=1)
    items = [
        "Certificate",
        "Declaration",
        "Acknowledgement",
        "Abstract / Executive Summary",
        "Chapter 1: Introduction",
        "Chapter 2: System Study",
        "Chapter 3: System Analysis and Design",
        "Chapter 4: UML and Data Design Diagrams",
        "Chapter 5: Implementation",
        "Chapter 6: Testing",
        "Chapter 7: Results and Discussion",
        "Chapter 8: Conclusion and Future Scope",
        "Chapter 9: References",
        "Chapter 10: Appendices",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
    doc.add_page_break()


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    for style_name in ["Normal", "Body Text"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.08

    for style_name, size in [("Title", 22), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Project Report")
    r.font.name = "Arial"
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("Optimizing Retail Stock Levels Through\nAI-Powered Demand Prediction")
    rt.font.name = "Arial"
    rt.font.size = Pt(24)
    rt.font.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = subtitle.add_run("Enhanced technical report with implementation evidence, figures, tables, and code excerpts")
    rs.font.name = "Arial"
    rs.font.size = Pt(12)
    rs.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_paragraph()
    meta = doc.add_table(rows=5, cols=2)
    meta_rows = [
        ("Student Name", "Mohit Yadav"),
        ("Enrollment No.", "024MSD110076"),
        ("Programme", "Master of Science (Data Science) MSc (DS)"),
        ("Guide / Mentor", "Vikas Kumar Atray"),
        ("Project Scope", "Store 1; BEVERAGES, DAIRY, and GROCERY I"),
    ]
    for row, (label, value) in zip(meta.rows, meta_rows):
        row.cells[0].text = label
        row.cells[1].text = value
    style_table(meta)
    doc.add_page_break()


def add_project_evidence(doc: Document) -> None:
    doc.add_heading("Implementation Evidence Added to This Version", level=2)
    add_note(
        doc,
        "Why this section is added",
        "This report version keeps the original written explanation but adds direct evidence from the project workspace: "
        "generated tables, figures, dashboard output, and source-code excerpts. This makes the report easier to verify.",
    )
    data = [
        ("Source / Output", "Purpose in report"),
        ("scripts/prepare_dataset.py", "Rebuilds the processed modelling dataset from raw Corporacion Favorita CSV files."),
        ("run_pipeline.py", "Runs forecasting, model comparison, inventory recommendations, simulation, and dashboard generation."),
        ("Reports/tables/model_metrics.csv", "Stores the model comparison values used in the results section."),
        ("Reports/tables/inventory_recommendations.csv", "Stores safety stock, reorder point, and recommended order quantity values."),
        ("Reports/tables/stock_alerts.csv", "Stores reorder alerts generated by the stock simulation."),
        ("Reports/figures/*.png", "Provides report-ready figures for sales history, forecast comparison, and inventory simulation."),
    ]
    table = doc.add_table(rows=len(data), cols=2)
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value
    style_table(table)
    doc.add_paragraph()


def add_dataset_summary(doc: Document) -> None:
    doc.add_heading("Dataset and Generated File Summary", level=3)
    data = [
        ("File", "Role"),
        ("train.csv", "Primary sales history used to build demand series."),
        ("items.csv", "Maps item numbers to product families."),
        ("stores.csv", "Provides store-level metadata."),
        ("transactions.csv", "Adds daily store activity information."),
        ("oil.csv", "Adds oil-price time-series context."),
        ("holidays_events.csv", "Adds calendar and event information."),
        ("store1_three_families_daily.csv", "Processed daily modelling dataset used by the pipeline."),
    ]
    table = doc.add_table(rows=len(data), cols=2)
    for r_idx, row in enumerate(data):
        for c_idx, value in enumerate(row):
            table.cell(r_idx, c_idx).text = value
    style_table(table)


def add_results_package(doc: Document) -> None:
    doc.add_heading("Generated Result Tables and Figures", level=2)
    add_csv_table(doc, TABLES / "model_metrics.csv", "Table: Forecasting Model Comparison", round_numeric=True)
    add_csv_table(doc, TABLES / "inventory_recommendations.csv", "Table: Inventory Recommendations", round_numeric=True)
    add_csv_table(doc, TABLES / "model_availability.csv", "Table: Model Availability")
    add_csv_table(doc, TABLES / "stock_alerts.csv", "Table: Sample Reorder Alerts", max_rows=8, round_numeric=True)
    add_picture(doc, FIGURES / "sales_history.png", "Figure: Sales History of Selected Families")
    add_picture(doc, FIGURES / "best_model_forecast.png", "Figure: Best Model Forecast vs Actual Sales")
    add_picture(doc, FIGURES / "inventory_simulation.png", "Figure: Inventory Simulation Output")


def add_design_figures(doc: Document) -> None:
    doc.add_heading("System and Design Diagrams", level=2)
    add_picture(doc, FIGURES / "figure_3_1_system_architecture.png", "Figure: System Architecture")
    add_picture(doc, FIGURES / "figure_3_2_dfd_level_0.png", "Figure: DFD Level 0")
    add_picture(doc, FIGURES / "figure_3_3_dfd_level_1.png", "Figure: DFD Level 1")
    add_picture(doc, FIGURES / "figure_4_1_use_case_diagram.png", "Figure: Use Case Diagram")
    add_picture(doc, FIGURES / "figure_4_2_activity_diagram.png", "Figure: Activity Diagram")
    add_picture(doc, FIGURES / "figure_4_3_sequence_diagram.png", "Figure: Sequence Diagram")
    add_picture(doc, FIGURES / "figure_4_4_data_schema.png", "Figure: Data Schema")


def add_code_appendix(doc: Document) -> None:
    doc.add_heading("Appendix E: Important Code Excerpts", level=1)
    add_code_block(
        doc,
        "Feature Engineering and Chronological Split",
        read_code_lines(SRC / "features.py", 8, 36),
    )
    add_code_block(
        doc,
        "Inventory Recommendation Formula",
        read_code_lines(SRC / "inventory.py", 8, 40),
    )
    add_code_block(
        doc,
        "Model Evaluation and Selection",
        read_code_lines(ROOT / "run_pipeline.py", 22, 74),
    )
    add_code_block(
        doc,
        "Pipeline Outputs",
        read_code_lines(ROOT / "run_pipeline.py", 76, 132),
    )


def copy_original_content_with_insertions(doc: Document) -> None:
    original = Document(str(ORIGINAL_DOCX))
    inserted = set()
    skip_tail_inserted_labels = False

    skip_exact = {
        "Table of Contents",
    }

    for paragraph in original.paragraphs:
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        if text in skip_exact:
            continue
        if text == "3.6 DFD and Use Case Diagrams":
            skip_tail_inserted_labels = True
            continue
        if skip_tail_inserted_labels and text == "Chapter 9: References":
            skip_tail_inserted_labels = False
        if skip_tail_inserted_labels:
            continue

        style_name = paragraph.style.name if paragraph.style is not None else "Normal"
        if style_name.startswith("Heading 1"):
            new_p = doc.add_heading(text, level=1)
        elif style_name.startswith("Heading 2"):
            new_p = doc.add_heading(text, level=2)
        elif style_name.startswith("Heading 3"):
            new_p = doc.add_heading(text, level=3)
        else:
            new_p = doc.add_paragraph()
            # Preserve the user's wording, but present it in the new report's body style.
            new_p.add_run(paragraph.text)

        if text == "1.11 Deliverables of the Project" and "evidence" not in inserted:
            add_project_evidence(doc)
            inserted.add("evidence")
        elif text == "2.8 Dataset Study" and "dataset" not in inserted:
            add_dataset_summary(doc)
            inserted.add("dataset")
        elif text == "3.5 System Architecture" and "design" not in inserted:
            add_design_figures(doc)
            inserted.add("design")
        elif text == "7.1 Forecasting Results" and "results" not in inserted:
            add_results_package(doc)
            inserted.add("results")

    if "results" not in inserted:
        doc.add_page_break()
        add_results_package(doc)
    doc.add_page_break()
    add_code_appendix(doc)


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "Optimizing Retail Stock Levels Through AI-Powered Demand Prediction"
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(100, 100, 100)

        footer = section.footer.paragraphs[0]
        footer.text = "Enhanced technical report"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(100, 100, 100)


def main() -> None:
    doc = Document()
    setup_styles(doc)
    add_cover(doc)
    add_static_contents(doc)
    copy_original_content_with_insertions(doc)
    add_header_footer(doc)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
